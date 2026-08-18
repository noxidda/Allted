import os
import io
import re

def convert_document(input_path_or_bytes, file_name, target_ext):
    clean_target = target_ext.lower().strip('.')
    source_ext = os.path.splitext(file_name)[1].lower().strip('.')

    if isinstance(input_path_or_bytes, bytes):
        input_bytes = input_path_or_bytes
    else:
        with open(input_path_or_bytes, 'rb') as f:
            input_bytes = f.read()

    extracted_text = _extract_text(input_bytes, source_ext, file_name)

    output_bytes = b''
    mime_type = 'text/plain'

    if clean_target == 'txt':
        output_bytes = extracted_text.encode('utf-8')
        mime_type = 'text/plain; charset=utf-8'
    elif clean_target == 'html':
        output_bytes = _format_to_html(extracted_text, file_name).encode('utf-8')
        mime_type = 'text/html; charset=utf-8'
    elif clean_target == 'md':
        output_bytes = _format_to_markdown(extracted_text).encode('utf-8')
        mime_type = 'text/markdown; charset=utf-8'
    elif clean_target == 'rtf':
        output_bytes = _format_to_rtf(extracted_text).encode('utf-8')
        mime_type = 'application/rtf'
    elif clean_target == 'pdf':
        output_bytes = _create_pdf_from_text(file_name, extracted_text)
        mime_type = 'application/pdf'
    elif clean_target in ['docx', 'doc', 'odt', 'epub']:
        output_bytes = _create_docx_from_text(file_name, extracted_text)
        mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    else:
        output_bytes = extracted_text.encode('utf-8')
        mime_type = 'text/plain; charset=utf-8'

    base_name = os.path.splitext(file_name)[0]
    out_file_name = f"{base_name}.{clean_target}"

    return {
        "bytes": output_bytes,
        "file_name": out_file_name,
        "mime_type": mime_type
    }

def _extract_text(input_bytes, source_ext, file_name):
    if source_ext == 'docx':
        try:
            import docx
            doc = docx.Document(io.BytesIO(input_bytes))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        paragraphs.append(row_text)
            if paragraphs:
                return "\n\n".join(paragraphs)
        except Exception as e:
            pass

    if sourceExt_is_pdf(source_ext, input_bytes):
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(input_bytes))
            text_pages = [page.extract_text() for page in reader.pages if page.extract_text()]
            if text_pages:
                return "\n\n".join(text_pages)
        except Exception as e:
            pass

    if source_ext in ['html', 'htm']:
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(input_bytes, 'html.parser')
            for script in soup(["script", "style"]):
                script.extract()
            text = soup.get_text(separator='\n')
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            if lines:
                return "\n".join(lines)
        except Exception as e:
            pass

    if source_ext == 'rtf':
        try:
            text = input_bytes.decode('utf-8', errors='ignore')
            text = re.sub(r'\\fonttbl[\s\S]*?\}', '', text)
            text = re.sub(r'\\colortbl[\s\S]*?\}', '', text)
            text = re.sub(r'\\[a-z]+\d*', '', text)
            text = re.sub(r'[\{\}]', '', text)
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            if lines:
                return "\n".join(lines)
        except Exception as e:
            pass

    try:
        text = input_bytes.decode('utf-8', errors='ignore')
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F]', '', text)
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        if lines:
            return "\n".join(lines)
    except Exception:
        pass

    return f"Content extracted from {file_name}"

def sourceExt_is_pdf(source_ext, input_bytes):
    return source_ext == 'pdf' or input_bytes.startswith(b'%PDF')

def _create_pdf_from_text(title, text):
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        buffer = io.BytesIO()
        pdf = canvas.Canvas(buffer, pagesize=letter)
        pdf.setTitle(title)

        pdf.setFont("Helvetica-Bold", 16)
        pdf.setFillColorRGB(0.48, 0.22, 0.92) # Purple header
        pdf.drawString(50, 750, title)

        pdf.setFont("Helvetica", 10)
        pdf.setFillColorRGB(0.12, 0.06, 0.21)

        y = 720
        lines = text.split('\n')
        for line in lines:
            if y < 50:
                pdf.showPage()
                pdf.setFont("Helvetica", 10)
                pdf.setFillColorRGB(0.12, 0.06, 0.21)
                y = 750

            wrapped_chunks = [line[i:i+90] for i in range(0, len(line), 90)] or [""]
            for chunk in wrapped_chunks:
                pdf.drawString(50, y, chunk)
                y -= 14
                if y < 50:
                    pdf.showPage()
                    pdf.setFont("Helvetica", 10)
                    pdf.setFillColorRGB(0.12, 0.06, 0.21)
                    y = 750

        pdf.save()
        return buffer.getvalue()
    except Exception as e:
        # Fallback simple PDF generator if reportlab is not installed yet
        return _fallback_pdf(title, text)

def _fallback_pdf(title, text):
    clean_text = text.replace('(', '\\(').replace(')', '\\)')
    lines = clean_text.split('\n')
    stream_content = f"BT /F1 12 Tf 50 750 Td ({title}) Tj 0 -20 Td "
    for line in lines[:50]:
        stream_content += f"({line[:80]}) Tj 0 -14 Td "
    stream_content += "ET"

    pdf_body = f"""%PDF-1.4
1 0 obj <</Type /Catalog /Pages 2 0 R>> endobj
2 0 obj <</Type /Pages /Kids [3 0 R] /Count 1>> endobj
3 0 obj <</Type /Page /Parent 2 0 R /Resources <</Font <</F1 4 0 R>>>> /Contents 5 0 R>> endobj
4 0 obj <</Type /Font /Subtype /Type1 /BaseFont /Helvetica>> endobj
5 0 obj <</Length {len(stream_content)}>>
stream
{stream_content}
endstream
endobj
xref
0 6
0000000000 65535 f 
0000000010 00000 n 
0000000060 00000 n 
0000000117 00000 n 
0000000223 00000 n 
0000000298 00000 n 
trailer <</Size 6 /Root 1 0 R>>
startxref
{350 + len(stream_content)}
%%EOF"""
    return pdf_body.encode('utf-8')

def _create_docx_from_text(title, text):
    try:
        import docx
        doc = docx.Document()
        doc.add_heading(title, level=1)
        for paragraph in text.split('\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    except Exception:
        # Simple HTML Word fallback
        html = f"<html><body><h1>{title}</h1>"
        for p in text.split('\n'):
            if p.strip():
                html += f"<p>{p.strip()}</p>"
        html += "</body></html>"
        return html.encode('utf-8')

def _format_to_html(text, title):
    paragraphs = [f"<p>{p.strip()}</p>" for p in text.split('\n') if p.strip()]
    body = "\n".join(paragraphs)
    return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{title}</title>
    <style>
        body {{ font-family: 'Segoe UI', sans-serif; line-height: 1.6; padding: 2rem; color: #1e1035; background: #faf5ff; }}
        h1 {{ color: #7c3aed; border-bottom: 2px solid #e9d5ff; }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    {body}
</body>
</html>"""

def _format_to_markdown(text):
    return "\n\n".join([p.strip() for p in text.split('\n') if p.strip()])

def _format_to_rtf(text):
    clean = text.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}').replace('\n', '\\par\n')
    return f"{{\\rtf1\\ansi\\deff0{{\\fonttbl{{\\f0 Segoe UI;}}}}\\f0\\fs24 {clean}}}"
